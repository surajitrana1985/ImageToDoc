import os

# OCR related library imports
import pytesseract as tess
from PIL import Image
tess.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Word Document related library imports
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement


image_directory = "images"
word_directory = "worddocs"
count = 0

for file in os.listdir(image_directory):
    img_file_path = os.path.join(image_directory, file)
    file_has_header = file.find("header") != -1
    document = Document()
    # set font
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    # style.font.size = Pt(14)
    count += 1
    if os.path.isfile(img_file_path):
        image_file = Image.open(img_file_path)
        text = tess.image_to_string(image_file)
        text_array = text.split("\n")
        output_filename = os.path.splitext(file)[0]
        # Now Add below children to root xml tree
        # create xml element using OxmlElement
        shd = OxmlElement('w:background')
        line_count = 0
        for line in text_array:
            if file_has_header == True and line_count == 0:
                style.font.size = Pt(16)
                para = document.add_paragraph(line)
                para.paragraph_format.line_spacing = 1.75
            else:
                style.font.size = Pt(14)
                para = document.add_paragraph(line)
                para.paragraph_format.line_spacing = 1.0
            line_count += 1

            # Add attributes to the xml element
            shd.set(qn('w:color'), '000000')  # black color
            shd.set(qn('w:themeColor'), 'text1')
            shd.set(qn('w:themeTint'), 'F2')

        # document.add_page_break()
        # Add background element at the start of Document.xml using below
        document.element.insert(0, shd)
        # creating output filename with extension
        doc_name = output_filename + ".docx"
        # Add displayBackgroundShape element to setting.xml
        shd1 = OxmlElement('w:displayBackgroundShape')
        document.settings.element.insert(0, shd1)

        # Writing the output word documents to disk
        document.save(os.path.join(word_directory, doc_name))
print("Done writing the output word documents")
